# -*- coding: utf-8 -*-
"""
Created on Fri Mar 29 08:21:01 2024
Class to hold all kind of tools

@author: kristina
"""


from PyQt5 import QtCore, QtGui, QtWidgets
from PyQt5.QtWidgets import  QLabel,  QMessageBox, QTableWidget, QTableWidgetItem, QTextBrowser
from PyQt5.QtGui import QImage

from docx import Document
from docx.shared import Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

import logging
from datetime import datetime, time
import numpy as np
import pandas as pd
import os
import cv2
from pathlib import Path
import xml.etree.ElementTree as ET
import traceback
from paddleocr import  draw_ocr
from paddleocr import PaddleOCR


class c_ocr:        
    def __init__(self):
        
        self.dfColumnNames = [
            'testString',
            'patternFound',
            'xmin',
            'ymin',
            'xmax',
            'ymax'
            ]
        self._df_RefStringData = pd.DataFrame(columns = self.dfColumnNames)
        print(self._df_RefStringData) 
               
        self.label_ampel = QLabel
        self.label_green = QLabel
        self.label_currentSelectedTemplate = QLabel
        self.lbl_previewOrderTemplateImage = QLabel   
        self.tableWidget_searchStrings = QTableWidget
        self.textBrowser_log = QTextBrowser
        self._bestellVorlagen_labes = ""
        self._ocrModel = PaddleOCR
         
    def testDevice(self, webcamNr):
       logging.info(self.dt() + 'testDevice() - we now check device: ' + str(webcamNr))
       cap = cv2.VideoCapture(webcamNr)         
       if cap is None or not cap.isOpened():
           print('Warning: unable to open video source: ' + str(webcamNr))
           logging.info(self.dt() + 'Warning: unable to open video source: ' + str(webcamNr))
           return False
       else:
           print('YEAA: this video source seem ok: ' + str(webcamNr))
           logging.info(self.dt() + 'YEAA: this video source seem ok: ' + str(webcamNr))
           return True 

    '''
    MsgBox von pyQt / https://doc.qt.io/qt-6/qmessagebox.html  / - OK  Cancel
    '''
    def msgBoxInfoOkCancel(self,txt,title):
        logging.info(self.dt() + 'msgBoxInfoOkCance()')
        msgBox = QMessageBox()
        msgBox.setIcon(QMessageBox.Information)           
        msgBox.setText(txt)
        msgBox.setWindowTitle(title)
        msgBox.setStandardButtons(QMessageBox.Ok | QMessageBox.Cancel)  
        v = msgBox.exec()
        return v
    
    def msgBoxYesCancel(self,txt,title):
        logging.info(self.dt() + 'msgBoxYes()')
        msgBox = QMessageBox()
        msgBox.setIcon(QMessageBox.Information)           
        msgBox.setText(txt)
        msgBox.setWindowTitle(title)
        msgBox.setStandardButtons(QMessageBox.Yes | QMessageBox.Cancel)
        return msgBox.exec()  
    
    def dt(self):
        return datetime.now().strftime("%Y-%m-%d %H:%M:%S") + ' - '

    # Define a function called date_diff_in_Seconds which calculates the difference in seconds between two datetime objects
    def dateDifSeconds(self):
        # Calculate the time difference between dt2 and dt1
        startDate = datetime.strptime('2024-06-17 01:00:00', '%Y-%m-%d %H:%M:%S')
        timedelta = datetime.now() - startDate
        # Return the total time difference in seconds
        return timedelta.days * 24 * 3600 + timedelta.seconds

              
    def changeOnlyHeightOfImage(self, image, newHeight):
        h,w,c = image.shape
        down_width = w
        down_height = newHeight
        down_points = (down_width, down_height)
        resized_image = cv2.resize(image, down_points, interpolation= cv2.INTER_LINEAR)
        return  resized_image

    #Fill the table with SearchStrings/patterns from the Dataframe with true and false
    def fillQtableWithDataframe(self):
        #We only need 2 column from Dataset in word
        df = self._df_RefStringData[['testString', 'patternFound']] 
        numrows = df.shape[0]  # Gives number of rows
        numcols = df.shape[1]  # Gives number of columns        
        try:
            # Set colums and rows in QTableWidget
            self.tableWidget_searchStrings.setRowCount(0)
            self.tableWidget_searchStrings.setColumnCount(numcols)
            self.tableWidget_searchStrings.setRowCount(numrows)
            #Tabellenbreite dynamisch anpassen
            header = self.tableWidget_searchStrings.horizontalHeader()
            header.setSectionResizeMode(QtWidgets.QHeaderView.ResizeToContents)       
            header.setSectionResizeMode(0, QtWidgets.QHeaderView.Stretch)
            
            df = df.reset_index()  # make sure indexes pair with number of rows

            for index, row in df.iterrows():
                print(row['testString'], row['patternFound'])
                self.tableWidget_searchStrings.setItem(index, 0, QTableWidgetItem(row['testString']))                                                                     
                self.tableWidget_searchStrings.item(index, 0).setBackground(QtGui.QColor(255,0,0)) #red
                self.tableWidget_searchStrings.setItem(index, 1, QTableWidgetItem(str(row['patternFound'])))
                if row['patternFound'] == False:
                    self.tableWidget_searchStrings.item(index, 0).setBackground(QtGui.QColor(255,0,0)) #red
                    self.tableWidget_searchStrings.item(index, 1).setBackground(QtGui.QColor(255,0,0)) #red
                else:
                    self.tableWidget_searchStrings.item(index, 0).setBackground(QtGui.QColor(50,205,50)) #green
                    self.tableWidget_searchStrings.item(index, 1).setBackground(QtGui.QColor(50,205,50)) #green
                                     
            print(df)
        except Exception as e:
            print(e)
            logging.error(traceback.format_exc())   
    
    
    #Iterate the bboxes of the EXAMPLE Bestell-Vorlage
    #zb kommen wir hier mit 2 boxen rein die in MakeSense definiert wurden
    def getReferenzListFromBBoxes( self, ocrModel, box, img ):
        self._df_RefStringData = self._df_RefStringData.iloc[0:0]  # clear df
        for b in box:
            xmin = b[0]
            ymin = b[1]
            xmax = b[2]
            ymax = b[3]
            h = ymax - ymin
            w = xmax - xmin
             
            #print(b)
            imgO = cv2.imread(img)
            imgB = imgO.copy()
            hImg, wImg,_  = imgB.shape

            imgCrop= imgB[ymin:ymin+h, xmin:xmin+w]
            #plt.imshow(imgCrop)
            #plt.show()  
            # ocr mit dem ROI-Image
            result = ocrModel.ocr(imgCrop, cls=True) # !!  P A D D L E  -  O C R    P A D D L E  -  O C R   
            result = result[0]       
            boxes = [line[0] for line in result]
            txts = [line[1][0] for line in result]
            scores = [line[1][1] for line in result]

            controlString = ' '.join(txts)
            cleanStr = (''.join(txts)).replace(" ", "")
            cleanStr = cleanStr.replace(".", "")
            s = controlString  + ' (' + cleanStr  + ')'
            #print(s)
            self.textBrowser_log.append('>>' + s) #Append text to the GUI
            newRow = {
                'testString': cleanStr,
                'patternFound': False,
                'xmin': xmin,
                'ymin': ymin,
                'xmax': xmax,
                'ymax': ymax
               }
           
            self._df_RefStringData.loc[len(self._df_RefStringData)] = newRow  # only use with a RangeIndex!        
        print ( self._df_RefStringData)
            
         

    #from Image/annotation we read teh pos and boundingBoxes
    def read_AnnotationContent(self, img_path, annotation_path):
        pth = Path(img_path)                #assets/train/images/bestellVorl1.JPG
        xml_path =annotation_path+ pth.stem +'.xml'   # assets/train/labels/bestellVorl1.xml
        
        img = cv2.imread( img_path)
        tree = ET.parse(xml_path)
        root = tree.getroot()
        list_with_all_boxes = []
        for boxes in root.iter('object'):
            filename = root.find('filename').text
            ymin, xmin, ymax, xmax = None, None, None, None
            ymin = int(boxes.find("bndbox/ymin").text)
            xmin = int(boxes.find("bndbox/xmin").text)
            ymax = int(boxes.find("bndbox/ymax").text)
            xmax = int(boxes.find("bndbox/xmax").text)
            list_with_single_boxes = [xmin, ymin, xmax, ymax]
            list_with_all_boxes.append(list_with_single_boxes)
            
            cv2.rectangle(img,(int(xmin),int(ymin)),(int(xmax),int(ymax)),(55,255,155),5)   
            #cv2.imwrite('protocolsPath/test.jpg', img) 
        return filename, list_with_all_boxes, img
    
    '''
    user selection of template ist all time in: self.label_currentSelectedTemplate
    This way we can reload the template any time
    ..also we refill the referenceBuffer - Important for quick changes in fast scan practice
    '''
    def reNewPreviewRefBufferAndList(self):
        path_of_image = self.label_currentSelectedTemplate.text()
        pixmap = QtGui.QPixmap(path_of_image)
        if not pixmap.isNull():
            pixmap = pixmap.scaledToWidth(400)
            self.lbl_previewOrderTemplateImage.setPixmap(pixmap)                    
            # split filename in parts and show in variables                  
            
            name, box, imgAnno = self.read_AnnotationContent(path_of_image, self._bestellVorlagen_labes)           
            self.getReferenzListFromBBoxes( self._ocrModel, box, path_of_image)                    
            self.fillQtableWithDataframe()
            
            logging.info(self.dt() + 'previewSelectedListObject() - objectFileName:' + name)  
            print(self._df_RefStringData)
                                          
    def setTrueFalseInRefTabl(self, ocrTxr):
        df = self._df_RefStringData.reset_index()  # make sure indexes pair with number of rows
        for index, row in df.iterrows():
            print(row['testString'], row['patternFound'])
            if row['testString'] in ocrTxr:
                #print(rs + '  in ' +  ocrTxr +  ' gefunden')
                self._df_RefStringData.loc[self._df_RefStringData.testString == row['testString'],'patternFound'] = True
                #print (self._df_RefStringData)
                
        
    #https://www.includehelp.com/python/check-if-all-values-in-dataframe-column-are-the-same.aspx
    def isAllFoundTrue(self):
        arr = self._df_RefStringData['patternFound'].to_numpy()           
        for elem in arr:            
            if elem == False:
                return False
        return True 
       
      
    #ToDo: soft! auch die halb gefundenen anzeigen
                            
    def isThisSameLabelAsReference(self,  txts):        
        # wenigstens irgendwo im ocr string MUSS die referenz stehen - sonst FALSE 
        txts = [elem for elem in txts if elem.strip()]
        txts = [word.replace('.','') for word in txts]
        txts = [word.replace(' ','') for word in txts]
        
        # markieren der gefundenen elemente on webcamview OVER ocr result ( not presize good )                      
        # https://stackoverflow.com/questions/72893442/paddle-ocr-boundingbox-format
        # OK ..aber mit versatz??      
        
        # we show the rectange over the originalTemplate searchStrings 
        imgTempl = cv2.imread(self.label_currentSelectedTemplate.text())  
           
        for txt in txts:            
            self.setTrueFalseInRefTabl(txt) # HIT?   

        if self.isAllFoundTrue():
            bFound = True
        else:
            bFound = False
            
        df = self._df_RefStringData.reset_index()  # make sure indexes pair with number of rows
        for index, row in df.iterrows():
            print(row['testString'], row['patternFound'], row['xmin'], row['ymin'], row['xmax'], row['ymax'] )
            xMin =  row['xmin']
            yMin =  row['ymin']
            xMax =  row['xmax']
            yMax =  row['ymax']
            if row['patternFound'] == True:
                cv2.rectangle(imgTempl,(xMin,yMin),(xMax,yMax), (50,205,50), 10)     # green
            else:
                cv2.rectangle(imgTempl,(xMin,yMin),(xMax,yMax), (0,0,0), 10)     # RED
            
        self.fillQtableWithDataframe()   
        #überarbeitetes Template zeigen # Bounding Box
        frame = cv2.cvtColor(imgTempl, cv2.COLOR_BGR2RGB)
        img = QImage(frame, frame.shape[1],frame.shape[0],frame.strides[0],QImage.Format_RGB888)        
        pixmap = QtGui.QPixmap.fromImage(img)
        pixmap = pixmap.scaledToWidth(400)
        self.lbl_previewOrderTemplateImage.setPixmap(pixmap)       
            
        cv2.imwrite('protocolsPath/templateWithBBoxes.jpg', imgTempl) # für den report später ..so oder so..mit oder ohne BB   
        
        if bFound == True:            
             print("\n**********  GEFUNDEN - LABEL ist OK !!!!!!!  ************* ")               
             pixmap = QtGui.QPixmap(':/image/images/green.JPG')
             pixmap = pixmap.scaled(300, 300, QtCore.Qt.KeepAspectRatio)
             self.label_ampel.setPixmap(pixmap)   
             return True            
        else: 
             pixmap = QtGui.QPixmap(':/image/images/red.JPG')
             pixmap = pixmap.scaled(300,300, QtCore.Qt.KeepAspectRatio)
             self.label_ampel.setPixmap(pixmap)   
             return False
        
    
    def pp_ocrOnSingleImageCV2(self, camScreenPosX, camScreenPosY, ocrModel, cvImg):
        try:
            resultOCR = ocrModel.ocr(cvImg, cls=True)   # OCR  OCR  OCR  OCR  OCR         
        except AttributeError:
            logging.error(traceback.format_exc())
            return cvImg           
        try:  
            #print(type(result))
            if resultOCR is None:
                logging.error(self.dt() + 'pp_ocrOnSingleImageCV2() - result is none!' )                
                return cvImg
            #Wenn nichts gefunden  oder so lange kein etikett in range ist           
            for idx in range(len(resultOCR)):
                res = resultOCR[idx]
                if res is None:  
                    self.reNewPreviewRefBufferAndList()  # bestellvorlage neu einlesen + refeDataframe
                    return cvImg
                for line in res:
                    #print(line)  
                    pass
               
            #Detection result anzeigen bzw BBoxen auswerten        
            result = resultOCR[0]       
            boxes = [line[0] for line in result]
            txts = [line[1][0] for line in result]
            scores = [line[1][1] for line in result]  
                
            # Specifying font path for draw_ocr method
            font = os.path.join('PaddleOCR', 'doc', 'fonts', 'latin.ttf')    
            # draw annotations on image
            im_2show= draw_ocr(cvImg, boxes, txts, scores, font_path=font) 
            hImg, wImg,_  = im_2show.shape
            
            # AMPEL - PAMPEL
            
            self.isThisSameLabelAsReference( txts )
            
        except Exception as e:
            logging.error(traceback.format_exc())
            print(e)     
            return cvImg
        return im_2show
   
    
    #https://stackoverflow.com/questions/26752856/python-docx-set-table-cell-background-and-text-color
    def setTableBGcolor( self, tbl, color, rowNr, colNr ):
        cell_xml_element = tbl.rows[rowNr].cells[colNr]._tc  # e.g tbl.rows[1].cells[0]
        #RETRIEVE THE TABLE CELL PROPERTIES
        table_cell_properties = cell_xml_element.get_or_add_tcPr()
        #CREATE SHADING OBJECT
        shade_obj = OxmlElement('w:shd')
        #SET THE SHADING OBJECT
        shade_obj.set(qn('w:fill'), color)    #  green "32CD32"
        #APPEND THE PROPERTIES TO THE TABLE CELL PROPERTIES
        table_cell_properties.append(shade_obj)
        return tbl
      
   
    # W O R D   protocol  /   https://python-docx.readthedocs.io/en/latest/
    '''
    protocolsPath = protocolsPath/  
    path_of_image e.g. templateImgFilePath = assets/train/images/emb_06A_01_ ed_292481.JPG 
    item.text()  e.g. templateFileName     = emb_06A_01_ ed_292481.JPG
    ocrTmpFile e.g. imgOCRresultPath       = images/tmp/imgPaddleOCR.png        
    wordFileName                           = emb_06A_01_ ed_292481.docx
    errFree                                = True - detection was OK   
    '''    
    
    def writeWordProtocol(self, protocolPath, templateImgFilePath,  wordFileName,  imgOCRresultPath):       
    
        print(self._df_RefStringData)
        try:
            document = Document()
            document.add_heading('ALI - Test Protocol', 0)            

            p = document.add_paragraph( self.dt() +  ' - ' + wordFileName )  # paragraph
            if self.isAllFoundTrue():
                p.add_run('  -- SUCCESS --').bold = True                                            # BOLD
            else:
                p.add_run('  -- FAILURE - NO SUCCESS --').bold = True 
            #p.add_run(' and some ')
            #p.add_run('italic.').italic = True
            
            # ------------- Table - Block -----------------------------------
            
            #We only need 2 column from Dataset in word
            df = self._df_RefStringData[['testString', 'patternFound']] 
            # add a table to the end and create a reference variable
            # extra row is so we can add the header row 
            t = document.add_table(df.shape[0]+1, df.shape[1]) # rows and columns
            # Add borders
            t.style = 'TableGrid'
            # add the header rows.
            for j in range(df.shape[-1]):
                t.cell(0,j).text = df.columns[j]
            # add the rest of the data frame
            for i in range(df.shape[0]):
                for j in range(df.shape[-1]):
                    t.cell(i+1,j).text = str(df.values[i,j])
            
            # table header color
            t = self.setTableBGcolor( t, 'A9A9A9', 0, 0 )  # gray   
            t = self.setTableBGcolor( t, 'A9A9A9', 0, 1 )  # gray 
            #Iterate Dataframe and COLOR some cells in wordTable
            df = self._df_RefStringData.reset_index()  # make sure indexes pair with number of rows
            for index, row in df.iterrows():
                print(row['testString'], row['patternFound'])
                if row['patternFound'] == True:
                    t = self.setTableBGcolor( t, '32CD32', index+1, 0 )  # green
                else:
                    t = self.setTableBGcolor( t, 'FF0000', index+1, 0 )  # red                    
            
           
            document.add_paragraph('\n' + templateImgFilePath)  # paragraph
            templateWithBBoxes = 'protocolsPath/templateWithBBoxes.jpg'            
            document.add_picture(templateWithBBoxes, width=Inches(4.25))
            
            document.add_paragraph('\nPaddle OCR Result ')  # paragraph
            document.add_picture( imgOCRresultPath , width=Inches(5.25))

            #document.add_page_break()
            wf = protocolPath + wordFileName   #   protocolsPath/ok/   40075_emb_06A_01_ed_292481.docx
            document.save(wf)
            
        except Exception as e:
            logging.error(traceback.format_exc())
            print(e)     
           
        